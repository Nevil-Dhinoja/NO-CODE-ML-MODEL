from fpdf import FPDF
from docx import Document

# === Your SRS content ===
srs_text = """
SOFTWARE REQUIREMENTS SPECIFICATION (SRS)

Project: No-Code Machine Learning Model Builder

1. Introduction

1.1 Purpose
The purpose of this project is to create a user-friendly no-code platform for building and deploying machine learning models. The tool will allow users to upload datasets, preprocess data, select ML models, train, evaluate, and export models without writing any code.

1.2 Scope
The application will focus on classification models initially and provide an interactive UI for data handling and model training. It will target beginners and non-technical users interested in machine learning.

1.3 Definitions, Acronyms, and Abbreviations
- ML: Machine Learning
- UI: User Interface
- CSV: Comma Separated Values

2. Overall Description

2.1 Product Perspective
This is a standalone web application built using Streamlit. It integrates common ML libraries like Scikit-learn and Pandas to provide backend processing.

2.2 Product Functions
- Dataset upload and preview
- Data preprocessing (handling missing values)
- Model selection among Logistic Regression, Decision Tree, and KNN
- Model training and evaluation
- Export model for later use

2.3 User Classes and Characteristics
- Beginner ML enthusiasts with no coding experience
- Data analysts looking for quick prototyping tools

2.4 Operating Environment
- Local desktop or cloud environment with Python installed
- Streamlit web interface accessed via browser

3. Specific Requirements

3.1 Functional Requirements

ID    Requirement                                      Priority
FR1   Allow users to upload CSV datasets              High
FR2   Display dataset preview with columns and rows  High
FR3   Enable dropping of columns with missing values Medium
FR4   Provide UI for selecting ML models              High
FR5   Train the selected ML model on the dataset      High
FR6   Display evaluation metrics (accuracy, etc.)    Medium
FR7   Export trained model as Pickle file             Medium

3.2 Non-functional Requirements
- User-friendly and responsive UI
- Modular and extensible architecture
- Quick response time for training small to medium datasets

4. External Interface Requirements

4.1 User Interfaces
- Streamlit-based UI with file uploader, checkboxes, dropdowns, and result display panels.

4.2 Hardware Interfaces
- No special hardware requirements; standard computer or cloud environment.

5. Other Requirements

- Documentation and user guide to assist non-technical users.
- Version control integration with GitHub.

6. Future Enhancements (Post MVP)

- Support for regression models and clustering.
- Model hyperparameter tuning UI.
- Deployment options on cloud services.

------------------------------
End of Document
"""

# === Your internship report content ===
content = {
    "Internship Details": """
Name: Nevil Dhinoja
Enrollment No: D24DCE147
Institute: CHARUSAT University
Department: Computer Engineering
Designation: Machine Learning Engineer Intern
Internship Duration: 18th May 2025 - 20th June 2025
Reporting Period: 18th May 2025 - 24th May 2025
Project Title: No-Code Machine Learning Model Builder
""",
    "Week 1 Summary": """
1. Orientation & Introduction
- Attended onboarding sessions and interacted with the Arocom team.
- Understood the company's work culture, ongoing projects, and client expectations.
- Discussed the scope and requirements of the internship project.

2. Research & Problem Understanding
- Conducted detailed research on the project: "No-Code Machine Learning Model Builder".
- Explored tools such as Streamlit, KNIME, Teachable Machine, and Hugging Face Spaces.
- Identified the gaps in current no-code tools and defined the unique value proposition of the project.

3. Project Initialization
- Finalized the architecture: modules for data upload, preprocessing, model selection, training, evaluation, and export.
- Set up the environment with Python, Streamlit, and Git for version control.
- Started prototyping basic UI with Streamlit and implemented initial layout structure.
- Created roadmap with milestone-based progress tracking.
""",
    "Practically Working Output": """
- Initial Streamlit App Created: Users can upload CSV datasets through a file uploader UI.
- Data Preview: The uploaded dataset is displayed in tabular format with column headers.
- Basic Statistics Module: Displays number of rows, columns, and missing value counts.
- Data Cleaning Preview: User can drop columns with missing values via checkbox controls.
- ML Model Selection (Stub): UI created for choosing between Logistic Regression, Decision Tree, or KNN (functionality to be added in Week 2).
""",
    "Skills & Tools Used": """
- Languages & Libraries: Python, Pandas, Scikit-learn, Streamlit
- Tools: GitHub (version control), Jupyter Notebook (experimentation), VS Code (IDE)
- Platforms: Google Colab (data testing), Streamlit local server
""",
    "Outcomes & Learnings": """
- Understood the UI design and functional workflow of a no-code ML tool.
- Implemented first working version of dataset ingestion and preprocessing UI.
- Learned modular development using Streamlit for quick deployment and testing.
""",
    "Next Week Goals": """
- Complete ML model selection logic and training output display.
- Implement train-test split functionality and evaluation metrics (accuracy, precision, recall).
- Add model download/export capability (Pickle format).
- Begin deployment testing on Streamlit Cloud (free hosting).
"""
}

# === Function to save SRS as Word document (.docx) ===
def save_srs_word(text, filename):
    doc = Document()
    for line in text.strip().split('\n'):
        line = line.strip()
        if line == "":
            doc.add_paragraph()
        elif line.isupper() and len(line.split()) < 10:
            # treat as heading if uppercase and short
            doc.add_heading(line, level=1)
        elif line.startswith("1.") or line.startswith("2.") or line.startswith("3.") or line[0].isdigit():
            # add as normal paragraph (could improve by detecting section numbering)
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)
    doc.save(filename)
    print(f"SRS saved successfully as {filename}")

# === Function to save Internship report as PDF ===
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, "Arocom IT Solutions Pvt. Ltd. - Internship Weekly Report", ln=True, align="C")
        self.ln(5)

    def chapter_title(self, title):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, title, ln=True, align="L")
        self.ln(2)

    def chapter_body(self, body):
        self.set_font("Arial", "", 11)
        self.multi_cell(0, 8, body)
        self.ln()

def save_internship_report_pdf(content_dict, filename):
    pdf = PDF()
    pdf.add_page()
    for title, body in content_dict.items():
        pdf.chapter_title(title)
        pdf.chapter_body(body)
    pdf.output(filename)
    print(f"Internship Report PDF saved successfully as {filename}")

# === Run the saving functions ===
save_srs_word(srs_text, "D24DCE147_Nevil_Dhinoja_SRS.docx")
save_internship_report_pdf(content, "D24DCE147_Nevil_Dhinoja_Week1_Internship_Report.pdf")
